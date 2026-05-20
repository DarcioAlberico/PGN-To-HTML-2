# -*- coding: utf-8 -*-
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk
import re
import os
import shutil
import sys
import ctypes
import threading
import webbrowser
import tempfile
import queue
import io
from datetime import datetime
from html.parser import HTMLParser
from typing import Callable

# Dependência principal para parsing robusto
import chess.pgn
import chess

# Gerador dedicado de diagramas
try:
    from . import chess_diagrams
    from .html_export import (
        CSS,
        build_css,
        ensure_diagram_font_css,
        gerar_html_final,
        get_css_preset_options,
        load_css_preset,
        write_html_bundle,
    )
    from .models import ConversionResult, DiagramAsset, ExerciseItem, GameSummary
    from .pdf_export import write_pdf_file
    from .pgn_validation import format_validation_report, validate_pgn
except ImportError:
    import chess_diagrams
    from html_export import (
        CSS,
        build_css,
        ensure_diagram_font_css,
        gerar_html_final,
        get_css_preset_options,
        load_css_preset,
        write_html_bundle,
    )
    from models import ConversionResult, DiagramAsset, ExerciseItem, GameSummary
    from pdf_export import write_pdf_file
    from pgn_validation import format_validation_report, validate_pgn

# ==================== EPUB (opcional) ====================
try:
    from ebooklib import epub
    HAS_EPUB = True
except ImportError:
    HAS_EPUB = False

# ==================== HTML VIEWER (opcional) ====================
try:
    os.environ.setdefault("QTWEBENGINE_DISABLE_SANDBOX", "1")
    from PySide6.QtCore import QEventLoop, QTimer, QUrl
    from PySide6.QtGui import QColor, QFont, QSyntaxHighlighter, QTextCharFormat, QTextCursor, QTextFormat, Qt
    from PySide6.QtWebEngineCore import QWebEngineSettings
    from PySide6.QtWidgets import (
        QApplication,
        QCheckBox,
        QComboBox,
        QFileDialog,
        QHBoxLayout,
        QLabel,
        QMainWindow,
        QMessageBox,
        QPlainTextEdit,
        QProgressBar,
        QPushButton,
        QSplitter,
        QTabWidget,
        QTextEdit,
        QVBoxLayout,
        QWidget,
    )
    from PySide6.QtWebEngineWidgets import QWebEngineView

    HAS_HTML_VIEWER = True
    HTML_VIEWER_IMPORT_ERROR = None
except ImportError:
    HAS_HTML_VIEWER = False
    HTML_VIEWER_IMPORT_ERROR = "PySide6.QtWebEngineWidgets"

# ==================== DOCX (opcional) ====================
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
    from docx.oxml.ns import qn

    HAS_DOCX = True
    DOCX_IMPORT_ERROR = None
except ImportError:
    HAS_DOCX = False
    DOCX_IMPORT_ERROR = "python-docx"

# ====================== FUNÇÕES AUXILIARES ======================

def _escape_html(s):
    if not s: return ""
    s = s.replace("&", "&amp;")
    s = s.replace("<", "&lt;")
    s = s.replace(">", "&gt;")
    return s

def _generate_analysis_links(fen):
    """Gera HTML dos botões de análise para uma FEN."""
    if not fen: return ""
    
    # Encode fen for URL (spaces to %20)
    import urllib.parse
    encoded_fen = urllib.parse.quote(fen)
    
    lichess_url = f"https://lichess.org/analysis/standard/{fen.replace(' ', '_')}"
    chesscom_url = f"https://www.chess.com/analysis?fen={encoded_fen}"
    
    return f'''
    <div class="analysis-links">
        <a href="{lichess_url}" target="_blank" class="lichess-btn">Lichess</a>
        <a href="{chesscom_url}" target="_blank" class="chesscom-btn">Chess.com</a>
    </div>
    '''


def _render_notice(message, tone="warning"):
    color_map = {
        "warning": ("#8a6d3b", "#fcf8e3", "#d6b656"),
        "error": ("#a94442", "#f2dede", "#d9534f"),
        "info": ("#31708f", "#d9edf7", "#5bc0de"),
    }
    text_color, background, border = color_map.get(tone, color_map["warning"])
    return (
        '<div style="margin:14px 0 22px; padding:10px 14px; '
        f'background:{background}; color:{text_color}; border-left:4px solid {border};'
        ' font-size:0.92em;">'
        f'{_escape_html(message)}</div>'
    )


def _clean_header_value(value):
    cleaned = (value or "").strip()
    if not cleaned:
        return ""
    if cleaned.lower() in {"unknown", "n/a", "na", "none", "null"}:
        return ""
    if re.fullmatch(r"[\?\.\-/:* ]+", cleaned):
        return ""
    return cleaned


def _build_header_html(idx, white, black, event, site, date, eco=""):
    clean_white = _clean_header_value(white)
    clean_black = _clean_header_value(black)
    clean_event = _clean_header_value(event)
    clean_site = _clean_header_value(site)
    clean_date = _clean_header_value(date)
    clean_eco = _clean_header_value(eco)

    if clean_white and clean_black:
        title_text = f"({idx}) {clean_white} – {clean_black}"
    elif clean_white:
        title_text = f"({idx}) {clean_white}"
    elif clean_black:
        title_text = f"({idx}) {clean_black}"
    else:
        title_text = f"({idx})"

    if clean_eco:
        title_text += f" [{clean_eco}]"

    info_parts = [part for part in (clean_event, clean_site, clean_date) if part]
    info_html = (
        f'<div class="info">{" • ".join(_escape_html(part) for part in info_parts)}</div>'
        if info_parts
        else ""
    )

    return f'<div class="headers"><div>{_escape_html(title_text)}</div>{info_html}</div>'


def build_game_summary(game, idx):
    headers = game.headers
    return GameSummary(
        index=idx,
        white=_clean_header_value(headers.get("White", "")),
        black=_clean_header_value(headers.get("Black", "")),
        event=_clean_header_value(headers.get("Event", "")),
        site=_clean_header_value(headers.get("Site", "")),
        date=_clean_header_value(headers.get("Date", "")),
        eco=_clean_header_value(headers.get("ECO", "")),
        result=_clean_header_value(headers.get("Result", "")),
        anchor=f"game-{idx}",
    )


def _is_parse_failure(exc):
    parse_error_types = (
        ValueError,
        AssertionError,
        chess.IllegalMoveError,
        chess.InvalidMoveError,
        chess.AmbiguousMoveError,
    )
    return isinstance(exc, parse_error_types)


DIAGRAM_COMMENT_PATTERN = re.compile(r"(?:#diagram\b|\[%\s*diagram\s*\])", re.IGNORECASE)
EXERCISE_COMMENT_PATTERN = re.compile(
    r"(?:#exercise\b(?P<hash_text>.*)|\[%\s*exercise(?:\s+\"(?P<bracket_quoted>[^\"]*)\"|\s+(?P<bracket_text>[^\]]*))?\])",
    re.IGNORECASE | re.DOTALL,
)


def comment_requests_diagram(comment):
    return bool(comment and DIAGRAM_COMMENT_PATTERN.search(comment))


def strip_diagram_marker(comment):
    cleaned = DIAGRAM_COMMENT_PATTERN.sub("", comment or "")
    cleaned = re.sub(r"\s{2,}", " ", cleaned)
    return cleaned.strip(" \t\r\n,;:-")


def parse_exercise_marker(comment):
    match = EXERCISE_COMMENT_PATTERN.search(comment or "")
    if not match:
        return None
    question = (
        match.group("hash_text")
        or match.group("bracket_quoted")
        or match.group("bracket_text")
        or ""
    )
    return re.sub(r"\s+", " ", question).strip(" \t\r\n,;:-\"")


def comment_requests_exercise(comment):
    return parse_exercise_marker(comment) is not None


def strip_exercise_marker(comment):
    cleaned = EXERCISE_COMMENT_PATTERN.sub("", comment or "")
    cleaned = re.sub(r"\s{2,}", " ", cleaned)
    return cleaned.strip(" \t\r\n,;:-")


def clean_control_markers(comment):
    return strip_exercise_marker(strip_diagram_marker(comment))

# Mapa de NAGs (Numeric Annotation Glyphs) para símbolos
NAG_MAP = {
    1: "!",
    2: "?",
    3: "‼",
    4: "⁇",
    5: "!?",
    6: "?!",
    7: "□", # Only move
    8: "□", # Only move
    11: "=",
    14: "⩲",
    15: "⩱",
    16: "±",
    17: "∓",
    18: "+-",
    19: "-+",
    22: "⨀", # Zugzwang
    23: "⨀", # Zugzwang
    26: "○",
    27: "○",
    32: "⟳", # Development
    33: "⟳", # Development
    36: "↑", 
    37: "↑", 
    40: "→", # Attack
    41: "→", 
    44: "⯹", # Compensation
    132: "⇆", # Counterplay
    133: "⇆", # Counterplay
    138: "⨁", # zeitnot
    139: "⨁", # zeitnot
    140: "∆", # With the idea...
    141: "∇", # Aimed against...
    142: "⌓", # Counterplay
}    

def get_nag_symbol(nag_int):
    return NAG_MAP.get(nag_int, "")

# ====================== WALKER DE JOGO (Visitor Pattern) ======================

class GameHtmlBuilder:
    def __init__(self, diagram_style=chess_diagrams.CLASSIC_DIAGRAM_STYLE):
        self.html_parts = []
        self.exercise_parts = []
        self.exercises = []
        self.diagram_assets = []
        self.warnings = []
        self.diagram_style = chess_diagrams.normalize_diagram_style(diagram_style)
        self.diagram_counter = 0
    
    def render_diagram_block(self, fen, output_dir, base_name, game_idx, warning_label):
        self.diagram_counter += 1
        diagram_folder = os.path.join(output_dir, "Diagrams")
        try:
            diagram_result = chess_diagrams.render_diagram_html(
                fen,
                output_dir=diagram_folder,
                base_name=base_name,
                idx=self.diagram_counter,
                size=360,
                style=self.diagram_style,
                web_root=output_dir,
            )
            asset = diagram_result.get("asset")
            if asset:
                self.diagram_assets.append(
                    DiagramAsset(
                        fen=fen,
                        svg_path=asset["svg_path"],
                        png_path=asset["png_path"],
                        web_path=asset["web_path"],
                    )
                )
            return diagram_result["html"], True
        except Exception as exc:
            self.warnings.append(
                f"Partida {game_idx}: nao foi possivel gerar {warning_label} ({exc})."
            )
            notice = (
                _render_notice(
                    f"Nao foi possivel gerar {warning_label} desta partida.",
                    tone="warning",
                )
            )
            return notice, False

    def append_diagram(self, fen, output_dir, base_name, game_idx, warning_label):
        html_block, ok = self.render_diagram_block(fen, output_dir, base_name, game_idx, warning_label)
        self.html_parts.append(html_block)
        if ok:
            self.html_parts.append(_generate_analysis_links(fen))
        return ok

    def append_exercise(self, fen, question, solution, output_dir, game_idx):
        exercise_index = len(self.exercises) + 1
        diagram_html, ok = self.render_diagram_block(
            fen,
            output_dir=output_dir,
            base_name=f"exercise_partida_{game_idx}",
            game_idx=game_idx,
            warning_label=f"o diagrama do exercicio {exercise_index}",
        )
        question_text = question or "Encontre o melhor lance."
        solution_text = solution or "Solucao nao informada."
        html_block = (
            f'<section class="exercise" id="exercise-{game_idx}-{exercise_index}">'
            f'<h2>Exercicio {exercise_index}</h2>'
            f'{diagram_html}'
            f'<p class="exercise-question">{_escape_html(question_text)}</p>'
            '<details class="exercise-solution">'
            '<summary>Solucao</summary>'
            f'<p>{_escape_html(solution_text)}</p>'
            '</details>'
            '</section>'
        )
        self.exercise_parts.append(html_block)
        self.exercises.append(
            ExerciseItem(
                game_index=game_idx,
                exercise_index=exercise_index,
                fen=fen,
                question=question_text,
                solution=solution_text,
                html=html_block,
            )
        )
        return ok


    def build_variation_string(self, node):
        """
        Constrói a string texto de uma variante completa, ex: "1. d4 (1. e4 e5) 1... d5"
        """
        tokens = []
        curr = node
        while True:
            # 1. O Lance
            board = curr.parent.board()
            san = board.san(curr.move)
            
            # Numeração
            if board.turn == chess.WHITE:
                num = f"{board.fullmove_number}."
            else:
                num = (f"{board.fullmove_number}..." if not tokens else "") # Só coloca reticências se for o inicio da string da variante? 
                # Melhor: sempre colocar número se for brancas, ou se for pretas e for o primeiro lance da variante.
                if not tokens: # Primeiro lance da variante
                     num = f"{board.fullmove_number}..."
                else:
                    num = "" # Em texto corrido de variante, lances pretos não costumam ter numero se seguem brancos
            
            # Ajuste fino: Se o lance anterior foi impresso, e agora é preto, não precisa numero.
            # Mas se acabamos de fechar uma sub-variante, talvez precise.
            # Simplificação: Sempre imprimir número para White. Para Black, imprimir se for o primeiro da seq.
            
            token_str = ""
            if board.turn == chess.WHITE:
                token_str = f"{board.fullmove_number}. {san}"
            else:
                 token_str = (f"{board.fullmove_number}... {san}" if not tokens else san)

            # NAGs
            for nag in curr.nags:
                token_str += get_nag_symbol(nag)

            tokens.append(token_str)
            
            # 2. Comentário Inline
            if curr.comment:
                cleaned_comment = clean_control_markers(curr.comment)
                if cleaned_comment:
                    c = _escape_html(cleaned_comment)
                    tokens.append(f'<span class="cmt">{c}</span>')

            # 3. Sub-variantes (aninhadas neste lance)
            if len(curr.variations) > 1:
                # variations[0] é a linha principal desta variante
                # variations[1...] são sub-variantes
                for subvar in curr.variations[1:]:
                    sub_text = self.build_variation_string(subvar)
                    tokens.append(f"({sub_text})")
            
            # 4. Avançar
            if not curr.variations:
                break
            curr = curr.variations[0] # Segue a main line desta variante
            
        return " ".join(tokens)


    def process_game(self, game, idx, output_dir="Diagrams"):
        self.html_parts = []
        self.exercise_parts = []
        self.exercises = []
        self.diagram_assets = []
        self.warnings = []
        self.diagram_counter = 0

        # --- Cabeçalho ---
        headers = game.headers
        white = headers.get("White", "?")
        black = headers.get("Black", "?")
        event = headers.get("Event", "?")
        site = headers.get("Site", "?")
        date = headers.get("Date", "????")
        eco = headers.get("ECO", "")
        fen = headers.get("FEN", "")
        setup = headers.get("SetUp", "0")

        header_html = _build_header_html(
            idx,
            white=white,
            black=black,
            event=event,
            site=site,
            date=date,
            eco=eco,
        )
        self.html_parts.append(header_html)

        # --- Diagrama Inicial ---
        if fen and setup == "1":
            self.append_diagram(
                fen,
                output_dir=output_dir,
                base_name=f"diagram_partida_{idx}",
                game_idx=idx,
                warning_label="o diagrama inicial",
            )

        # --- Comentário inicial (antes do primeiro lance) ---
        if game.comment:
            cleaned_game_comment = clean_control_markers(game.comment)
            if cleaned_game_comment:
                c = _escape_html(cleaned_game_comment)
                self.html_parts.append(f'<p class="comment">{c}</p>')

        # --- Iterar movimentos da Mainline ---
        node = game
        current_mainline_buffer = []

        while node.variations:
            # Pega o próximo lance da linha principal
            next_node = node.variations[0]
            
            # -- Variantes Alternativas (Nível 1) --
            # Se houver outras variantes além da main line neste ponto
            if len(node.variations) > 1:
                # Flush do buffer mainline antes de imprimir variantes
                if current_mainline_buffer:
                    txt = " ".join(current_mainline_buffer)
                    self.html_parts.append(f'<p class="mainline">{txt}</p>')
                    current_mainline_buffer = []
                
                # Processa variantes
                for variant_node in node.variations[1:]:
                    var_str = self.build_variation_string(variant_node)
                    self.html_parts.append(f'<p class="variant">({var_str})</p>')

            # -- Processa o lance Principal --
            board = node.board()
            san = board.san(next_node.move)
            
            # Formatação do lance (1. e4 ou 1... e5 se for o caso?)
            # Na mainline, sempre usamos numeração padrão
            move_str = ""
            if board.turn == chess.WHITE:
                move_str = f"{board.fullmove_number}. {san}"
            else:
                # Se for lance preto, verificamos se precisamos de "1... " 
                # (ex: começo de linha ou após variante/comentário)
                # Por simplicidade, na mainline: "1. e4 e5 2. Nf3"
                if not current_mainline_buffer: # Inicio de paragrafo
                     move_str = f"{board.fullmove_number}... {san}"
                else:
                     move_str = san

            # NAGs
            for nag in next_node.nags:
                move_str += get_nag_symbol(nag)
            
            current_mainline_buffer.append(move_str)
            
            # -- Comentários do Lance --
            if next_node.comment:
                wants_diagram = comment_requests_diagram(next_node.comment)
                exercise_question = parse_exercise_marker(next_node.comment)
                cleaned_comment = clean_control_markers(next_node.comment)
                # Flush mainline para mostrar comentário em parágrafo separado
                if current_mainline_buffer:
                    txt = " ".join(current_mainline_buffer)
                    self.html_parts.append(f'<p class="mainline">{txt}</p>')
                    current_mainline_buffer = []

                if cleaned_comment:
                    cmt = _escape_html(cleaned_comment)
                    self.html_parts.append(f'<p class="comment">{cmt}</p>')

                if wants_diagram:
                    self.append_diagram(
                        next_node.board().fen(),
                        output_dir=output_dir,
                        base_name=f"diagram_partida_{idx}",
                        game_idx=idx,
                        warning_label=f"o diagrama apos o lance {board.fullmove_number}",
                    )

                if exercise_question is not None:
                    solution = ""
                    if next_node.variations:
                        solution_board = next_node.board()
                        solution = solution_board.san(next_node.variations[0].move)
                    self.append_exercise(
                        next_node.board().fen(),
                        exercise_question,
                        solution,
                        output_dir=output_dir,
                        game_idx=idx,
                    )
                
                # Após comentário, se o próximo for Black, vai precisar de "N..."
                # Mas meu loop reinicia e verifica `current_mainline_buffer` vazio, então já trata.

            # Avançar
            node = next_node

        # Flush final
        if current_mainline_buffer:
            txt = " ".join(current_mainline_buffer)
            self.html_parts.append(f'<p class="mainline">{txt}</p>')
            
        # Resultado
        res = headers.get("Result", "*")
        if res != "*":
             self.html_parts.append(f'<p class="mainline" style="font-weight:bold; text-align:right">{res}</p>')

        return f'<section id="game-{idx}" class="game">\n' + "\n".join(self.html_parts) + "\n</section>"


# ====================== PARSER PRINCIPAL ======================

# ====================== FALLBACK PARSER (Regex Simples) ======================
# Usado quando o python-chess falha (ex: lances ilegais)

def fallback_parse_game(
    pgn_text,
    idx,
    output_dir=".",
    compatibility_note=None,
    diagram_style=chess_diagrams.CLASSIC_DIAGRAM_STYLE,
):
    """
    Parser simplificado baseado em regex para partidas inválidas/ilegais.
    Não valida lances, apenas formata o texto.
    """
    # 1. Extrair tags
    tags = dict(re.findall(r'\[(\w+)\s+"([^"]*)"\]', pgn_text))
    white = tags.get("White", "?")
    black = tags.get("Black", "?")
    event = tags.get("Event", "?")
    site = tags.get("Site", "?")
    date = tags.get("Date", "????")
    eco = tags.get("ECO", "")
    fen = tags.get("FEN", "")
    setup = tags.get("SetUp", "0")

    header_html = _build_header_html(
        idx,
        white=white,
        black=black,
        event=event,
        site=site,
        date=date,
        eco=eco,
    )

    if compatibility_note:
        header_html += _render_notice(compatibility_note, tone="warning")

    # 2. Diagrama (se houver FEN)
    diag_html = ""
    if fen and setup == "1":
        diagram_folder = os.path.join(output_dir, "Diagrams")
        try:
            diagram_result = chess_diagrams.render_diagram_html(
                fen,
                output_dir=diagram_folder,
                base_name=f"diagram_partida_{idx}_fallback",
                idx=1,
                style=diagram_style,
                web_root=output_dir,
            )
            diag_html = diagram_result["html"]
            # Links analise
            diag_html += _generate_analysis_links(fen)
        except Exception:
            pass # Se falhar diagrama no fallback, ignora

    # 3. Processar Corpo (Movimentos)
    body = re.sub(r'\[.*?\]\s*', '', pgn_text, flags=re.S) # Remove headers
    body = re.sub(r'\{.*?\}', lambda m: f' <span class="cmt">{_escape_html(m.group(0)[1:-1])}</span> ', body, flags=re.S) # Comentários
    body = re.sub(r'\((.*?)\)', lambda m: f' <p class="variant">({_escape_html(m.group(1))})</p> ', body, flags=re.S) # Variantes (simples)
    
    # Formatação básica de mainline
    lines = []
    for line in body.split('\n'):
        line = line.strip()
        if not line: continue
        lines.append(f'<p class="mainline">{line}</p>')
    
    return header_html + diag_html + "\n".join(lines)



# ====================== PARSER PRINCIPAL ======================

def _is_empty_default_game(game):
    if game is None:
        return True
    if game.comment:
        return False
    if list(game.mainline_moves()):
        return False

    default_headers = {
        "Event": "?",
        "Site": "?",
        "Date": "????.??.??",
        "Round": "?",
        "White": "?",
        "Black": "?",
        "Result": "*",
    }
    return all(game.headers.get(key) == value for key, value in default_headers.items())


def iter_pgn_games(pgn_text):
    """
    Itera partidas usando o parser do python-chess.

    Isso evita depender de uma tag [Event] no inicio de cada partida e aceita
    PGNs com tags em ordem incomum.
    """
    pgn_io = io.StringIO((pgn_text or "").lstrip("\ufeff"))
    while True:
        game = chess.pgn.read_game(pgn_io)
        if game is None:
            break
        if _is_empty_default_game(game):
            continue
        yield game


def _game_to_pgn_text(game):
    try:
        exporter = chess.pgn.StringExporter(headers=True, variations=True, comments=True)
        return game.accept(exporter)
    except Exception:
        return ""


def scan_pgn_headers(pgn_text):
    summaries = []
    for idx, game in enumerate(iter_pgn_games(pgn_text), start=1):
        summaries.append(build_game_summary(game, idx))
    return summaries


def filter_game_summaries(summaries, player="", eco="", result=""):
    player_filter = (player or "").strip().lower()
    eco_filter = (eco or "").strip().lower()
    result_filter = (result or "").strip().lower()

    filtered = []
    for summary in summaries:
        if player_filter:
            player_blob = " ".join([summary.white, summary.black]).lower()
            if player_filter not in player_blob:
                continue
        if eco_filter and eco_filter not in (summary.eco or "").lower():
            continue
        if result_filter and result_filter != (summary.result or "").lower():
            continue
        filtered.append(summary)
    return filtered


def convert_pgn(
    pgn_text,
    output_dir=".",
    progress_callback: Callable[[str], None] | None = None,
    diagram_style=chess_diagrams.CLASSIC_DIAGRAM_STYLE,
    selected_game_indexes=None,
    exercise_mode="book",
):
    """
    Converte um texto PGN em HTML e ativos associados.
    """
    indexed_games = list(enumerate(iter_pgn_games(pgn_text), start=1))
    if selected_game_indexes is not None:
        selected = {int(index) for index in selected_game_indexes}
        indexed_games = [(idx, game) for idx, game in indexed_games if idx in selected]

    normalized_style = chess_diagrams.normalize_diagram_style(diagram_style)
    result = ConversionResult(diagram_style=normalized_style)
    normalized_exercise_mode = (exercise_mode or "book").strip().lower()
    if normalized_exercise_mode not in {"book", "exercises", "both"}:
        raise ValueError(f"Modo de exercicio desconhecido: {exercise_mode}")
    total = len(indexed_games)

    if total == 0:
        if selected_game_indexes is not None:
            result.warnings.append("Nenhuma partida foi selecionada para conversao.")
            return result
        if pgn_text and pgn_text.strip():
            result.warnings.append(
                "Nenhuma partida PGN valida foi encontrada; tentando modo de compatibilidade."
            )
            result.blocks.append(
                fallback_parse_game(
                    pgn_text,
                    1,
                    output_dir=output_dir,
                    compatibility_note=(
                        "O texto nao foi reconhecido como PGN valido e foi "
                        "processado em modo de compatibilidade."
                    ),
                    diagram_style=normalized_style,
                )
            )
        return result

    for processed_count, (count, game) in enumerate(indexed_games, start=1):
        if progress_callback:
            progress_callback(f"Processando partida {processed_count}/{total}...")

        try:
            game_errors = getattr(game, "errors", None) or []
            for game_error in game_errors:
                result.warnings.append(f"Partida {count}: erro de PGN: {game_error}")

            builder = GameHtmlBuilder(diagram_style=normalized_style)
            html_block = builder.process_game(game, count, output_dir=output_dir)
            result.summaries.append(build_game_summary(game, count))
            result.exercise_blocks.extend(builder.exercise_parts)
            result.exercises.extend(builder.exercises)
            if normalized_exercise_mode == "book":
                result.blocks.append(html_block)
            elif normalized_exercise_mode == "exercises":
                result.blocks.extend(builder.exercise_parts)
            else:
                result.blocks.append(html_block)
                result.blocks.extend(builder.exercise_parts)
            result.diagram_assets.extend(builder.diagram_assets)
            result.warnings.extend(builder.warnings)

        except Exception as exc:
            if _is_parse_failure(exc):
                compatibility_note = (
                    "Esta partida contem lances ilegais ou nao padrao e foi "
                    "processada em modo de compatibilidade."
                )
            else:
                compatibility_note = (
                    "Esta partida foi processada em modo de compatibilidade por "
                    "causa de um erro interno durante o processamento estrito."
                )
                result.warnings.append(
                    f"Partida {count}: erro interno no modo estrito ({exc})."
                )

            try:
                raw_game = _game_to_pgn_text(game)
                html_fallback = fallback_parse_game(
                    raw_game,
                    count,
                    output_dir=output_dir,
                    compatibility_note=compatibility_note,
                    diagram_style=normalized_style,
                )
                result.blocks.append(html_fallback)
            except Exception as fallback_exc:
                raise RuntimeError(
                    f"Falha ao processar a partida {count}: {fallback_exc}"
                ) from fallback_exc

    return result


def processar_pgn_worker(
    pgn_text,
    progress_queue,
    output_dir=".",
    diagram_style=chess_diagrams.CLASSIC_DIAGRAM_STYLE,
    selected_game_indexes=None,
    exercise_mode="book",
):
    def _progress(message):
        progress_queue.put(("status", message))

    try:
        result = convert_pgn(
            pgn_text,
            output_dir=output_dir,
            progress_callback=_progress,
            diagram_style=diagram_style,
            selected_game_indexes=selected_game_indexes,
            exercise_mode=exercise_mode,
        )
        progress_queue.put(("done", result))
    except Exception as exc:
        progress_queue.put(("error", str(exc)))

def _clean_docx_text(text):
    if not text:
        return ""
    return re.sub(r"\s+", " ", text)


class HtmlPreviewRenderer:
    def __init__(self, project_dir):
        self.project_dir = os.path.abspath(project_dir)
        self.qt_app = QApplication.instance() or QApplication([])
        self.view = _ClickAwareWebEngineView(self._handle_mouse_click)
        self.view.setAttribute(Qt.WidgetAttribute.WA_NativeWindow, True)
        web_settings = self.view.settings()
        web_settings.setAttribute(QWebEngineSettings.WebAttribute.LocalContentCanAccessFileUrls, True)
        self.view.resize(920, 1200)
        self.view.loadFinished.connect(self._on_load_finished)
        self._line_click_callback = None
        self._load_finished_callback = None
        self._attached_hwnd = None
        self._current_base_dir = self.project_dir

    def attach_to_tk(self, tk_widget, line_click_callback=None, load_finished_callback=None):
        if sys.platform != "win32":
            raise RuntimeError("Viewer embutido disponivel apenas no Windows.")

        self._line_click_callback = line_click_callback
        self._load_finished_callback = load_finished_callback
        tk_widget.update_idletasks()

        parent_hwnd = int(tk_widget.winfo_id())
        view_hwnd = int(self.view.winId())

        user32 = ctypes.windll.user32
        style = user32.GetWindowLongPtrW(view_hwnd, -16)
        style = (style & ~0x80000000) | 0x40000000 | 0x10000000
        user32.SetWindowLongPtrW(view_hwnd, -16, style)
        user32.SetParent(view_hwnd, parent_hwnd)
        self._attached_hwnd = view_hwnd
        self.view.show()
        self.resize(tk_widget.winfo_width(), tk_widget.winfo_height())

    def resize(self, width, height):
        if self._attached_hwnd and sys.platform == "win32":
            ctypes.windll.user32.MoveWindow(
                self._attached_hwnd,
                0,
                0,
                max(1, int(width)),
                max(1, int(height)),
                True,
            )
        else:
            self.view.resize(max(1, int(width)), max(1, int(height)))

    def load_html(self, html_text, base_dir=None):
        self._current_base_dir = os.path.abspath(base_dir or self.project_dir)
        prepared_html = self._prepare_html(html_text)
        self.view.setHtml(prepared_html, QUrl.fromLocalFile(os.path.join(self._current_base_dir, "")))

    def load_html_document(self, html_text, html_path):
        self._current_base_dir = os.path.abspath(os.path.dirname(html_path) or self.project_dir)
        prepared_html = self._prepare_html(html_text)
        with open(html_path, "w", encoding="utf-8") as file_obj:
            file_obj.write(prepared_html)
        self.view.setUrl(QUrl.fromLocalFile(os.path.abspath(html_path)))

    def clear(self):
        self.view.setHtml("<!DOCTYPE html><html><body></body></html>", QUrl())

    def _prepare_html(self, html_text):
        return self._inject_source_line_markers(html_text)

    def _inject_source_line_markers(self, html_text):
        marked_lines = []
        tag_pattern = re.compile(r"<([A-Za-z][\w:-]*)(?=[\s>/])")

        for line_number, line in enumerate(html_text.splitlines(), start=1):
            def add_marker(match):
                tag_name = match.group(1).lower()
                if tag_name in {"html", "head", "body", "meta", "link", "style", "script", "title"}:
                    return match.group(0)
                return f'<{match.group(1)} data-source-line="{line_number}"'

            marked_lines.append(tag_pattern.sub(add_marker, line))

        return "\n".join(marked_lines)

    def _source_line_script(self, x, y):
        return f"""
(() => {{
    const x = {int(x)};
    const y = {int(y)};
    let el = document.elementFromPoint(x, y);
    if (!el) return null;

    let tagged = el.closest('[data-source-line]');
    if (!tagged) {{
        const elements = document.querySelectorAll('[data-source-line]');
        let best = null;
        let bestDistance = Infinity;
        for (const candidate of elements) {{
            const rect = candidate.getBoundingClientRect();
            const dy = Math.abs(rect.top - y);
            if (dy < bestDistance) {{
                best = candidate;
                bestDistance = dy;
            }}
        }}
        tagged = best;
    }}

    if (!tagged) return null;
    return parseInt(tagged.getAttribute('data-source-line'), 10) || null;
}})();
"""

    def _handle_mouse_click(self, x, y):
        if self._line_click_callback is None:
            return
        self.view.page().runJavaScript(
            self._source_line_script(x, y),
            self._on_source_line_result,
        )

    def _on_source_line_result(self, result):
        if result is None or self._line_click_callback is None:
            return
        try:
            line_number = int(result)
        except (TypeError, ValueError):
            return
        self._line_click_callback(line_number)

    def _on_load_finished(self, ok):
        if self._load_finished_callback is not None:
            self._load_finished_callback(bool(ok))

    def process_events(self):
        self.qt_app.processEvents()

    def get_source_line_from_point(self, x, y):
        script = f"""
(() => {{
    const x = {int(x)};
    const y = {int(y)};
    let el = document.elementFromPoint(x, y);
    if (!el) return null;

    let tagged = el.closest('[data-source-line]');
    if (!tagged) {{
        const elements = document.querySelectorAll('[data-source-line]');
        let best = null;
        let bestDistance = Infinity;
        for (const candidate of elements) {{
            const rect = candidate.getBoundingClientRect();
            const dy = Math.abs(rect.top - y);
            if (dy < bestDistance) {{
                best = candidate;
                bestDistance = dy;
            }}
        }}
        tagged = best;
    }}

    if (!tagged) return null;
    return parseInt(tagged.getAttribute('data-source-line'), 10) || null;
}})();
"""
        result = self._run_javascript(script)
        return int(result) if result else None

    def _document_metric(self, script):
        return self._run_javascript(script, timeout_ms=5000)

    def _run_javascript(self, script, timeout_ms=5000):
        loop = QEventLoop()
        state = {"done": False, "value": None}

        def on_result(value):
            if state["done"]:
                return
            state["done"] = True
            state["value"] = value
            loop.quit()

        self.view.page().runJavaScript(script, on_result)
        QTimer.singleShot(timeout_ms, loop.quit)
        loop.exec()
        return state["value"]

    def _wait(self, timeout_ms):
        loop = QEventLoop()
        QTimer.singleShot(timeout_ms, loop.quit)
        loop.exec()


class _ClickAwareWebEngineView(QWebEngineView):
    def __init__(self, click_callback):
        super().__init__()
        self._click_callback = click_callback

    def mousePressEvent(self, event):
        if self._click_callback and event.button() == Qt.MouseButton.LeftButton:
            try:
                position = event.position()
                self._click_callback(int(position.x()), int(position.y()))
            except Exception:
                pass
        super().mousePressEvent(event)


class DocxBlockParser(HTMLParser):
    def __init__(self, document, asset_lookup, temp_dir, diagram_style):
        super().__init__(convert_charrefs=True)
        self.document = document
        self.asset_lookup = asset_lookup
        self.temp_dir = temp_dir
        self.diagram_font = chess_diagrams.get_diagram_font_spec(diagram_style)
        self.tag_stack = []
        self.current_paragraph = None
        self.current_kind = None
        self.current_link_href = None
        self.current_link_parts = []
        self.in_comment_span = False
        self.capture_pre = False
        self.pre_parts = []
        self.pre_index = 0

    def handle_starttag(self, tag, attrs):
        attrs_dict = dict(attrs)
        classes = set(attrs_dict.get("class", "").split())
        self.tag_stack.append((tag, classes, attrs_dict))

        if tag == "p":
            if "mainline" in classes:
                self._start_paragraph("mainline")
            elif "comment" in classes:
                self._start_paragraph("comment")
            elif "variant" in classes:
                self._start_paragraph("variant")
            else:
                self._start_paragraph("body")
            return

        if tag == "div" and "headers" in classes:
            return

        if tag == "div" and self._inside_class("headers"):
            if "info" in classes:
                self._start_paragraph("header_info")
            else:
                self._start_paragraph("header_title")
            return

        if tag == "div" and ("analysis-links" in classes or "chess-ocr-links" in classes):
            self._start_paragraph("links")
            return

        if tag == "div" and "diagram" in classes:
            return

        if tag == "div" and classes:
            self._start_paragraph("notice")
            return

        if tag == "span" and "cmt" in classes:
            self.in_comment_span = True
            return

        if tag == "a":
            self.current_link_href = attrs_dict.get("href", "")
            self.current_link_parts = []
            return

        if tag == "img":
            self._insert_image_from_src(attrs_dict.get("src", ""))
            return

        if tag == "pre" and "chess-merida-diagram" in classes:
            self.capture_pre = True
            self.pre_parts = []
            return

        if tag == "br":
            if self.capture_pre:
                self.pre_parts.append("\n")
            elif self.current_paragraph is not None:
                self.current_paragraph.add_run("\n")

    def handle_endtag(self, tag):
        entry = self._pop_tag(tag)
        classes = entry[1] if entry else set()

        if tag == "a":
            self._flush_link()
            return

        if tag == "span" and "cmt" in classes:
            self.in_comment_span = False
            return

        if tag == "pre" and "chess-merida-diagram" in classes:
            self._insert_merida_pre()
            return

        if tag == "p":
            self._end_paragraph()
            return

        if tag == "div" and (
            "info" in classes
            or "analysis-links" in classes
            or "chess-ocr-links" in classes
            or self.current_kind in {"header_title", "header_info", "notice"}
        ):
            self._end_paragraph()

    def handle_data(self, data):
        if self.capture_pre:
            self.pre_parts.append(data)
            return

        if self.current_link_href is not None:
            self.current_link_parts.append(data)
            return

        if self.current_paragraph is None:
            return

        cleaned = _clean_docx_text(data)
        if not cleaned.strip():
            if self.current_paragraph.text and not self.current_paragraph.text.endswith(" "):
                self.current_paragraph.add_run(" ")
            return

        if not self.current_paragraph.text:
            cleaned = cleaned.lstrip()

        run = self.current_paragraph.add_run(cleaned)
        if self.current_kind in {"comment", "variant"} or self.in_comment_span:
            run.italic = True
        if self.current_kind == "header_title":
            run.bold = True

    def _start_paragraph(self, kind):
        self.current_paragraph = self.document.add_paragraph()
        self.current_kind = kind
        fmt = self.current_paragraph.paragraph_format

        if kind == "header_title":
            self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fmt.space_before = Pt(18)
            fmt.space_after = Pt(2)
        elif kind == "header_info":
            self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fmt.space_after = Pt(10)
        elif kind == "comment":
            fmt.left_indent = Inches(0.25)
            fmt.space_before = Pt(6)
            fmt.space_after = Pt(6)
        elif kind == "variant":
            fmt.left_indent = Inches(0.35)
            fmt.space_before = Pt(4)
            fmt.space_after = Pt(4)
        elif kind == "links":
            self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fmt.space_after = Pt(10)
        elif kind == "notice":
            fmt.left_indent = Inches(0.15)
            fmt.space_before = Pt(4)
            fmt.space_after = Pt(6)
        else:
            fmt.space_after = Pt(4)

    def _end_paragraph(self):
        self.current_paragraph = None
        self.current_kind = None

    def _inside_class(self, css_class):
        return any(css_class in classes for _, classes, _ in self.tag_stack)

    def _pop_tag(self, tag):
        for index in range(len(self.tag_stack) - 1, -1, -1):
            if self.tag_stack[index][0] == tag:
                entry = self.tag_stack[index]
                del self.tag_stack[index:]
                return entry
        return None

    def _flush_link(self):
        if self.current_paragraph is None:
            self._start_paragraph("links")

        label = _clean_docx_text("".join(self.current_link_parts)).strip()
        href = (self.current_link_href or "").strip()
        if self.current_paragraph.text:
            self.current_paragraph.add_run(" | ")

        if href and label and href != label:
            text = f"{label}: {href}"
        else:
            text = label or href

        run = self.current_paragraph.add_run(text)
        run.underline = True
        self.current_link_href = None
        self.current_link_parts = []

    def _insert_image_from_src(self, src):
        if not src:
            return

        normalized_src = src.replace("\\", "/")
        png_path = self.asset_lookup.get(normalized_src)
        if png_path is None:
            png_path = self.asset_lookup.get(os.path.basename(normalized_src))

        if not png_path or not os.path.isfile(png_path):
            return

        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(png_path, width=Inches(4.6))

    def _insert_merida_pre(self):
        diagram_text = "".join(self.pre_parts).strip("\n")
        self.capture_pre = False
        self.pre_parts = []

        if not diagram_text:
            return

        font_name = (self.diagram_font or {}).get("font_family") or "Chess Merida"
        font_path = (self.diagram_font or {}).get("font_path")

        try:
            paragraph = self.document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fmt = paragraph.paragraph_format
            fmt.space_before = Pt(6)
            fmt.space_after = Pt(6)
            fmt.line_spacing = 1.0

            run = paragraph.add_run(diagram_text)
            self._set_run_font(run, font_name, 28)
        except Exception:
            try:
                png_path = os.path.join(self.temp_dir, f"merida_docx_{self.pre_index}.png")
                self.pre_index += 1
                chess_diagrams.render_merida_diagram_png(
                    diagram_text,
                    png_path,
                    font_path=font_path,
                )
                paragraph = self.document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(png_path, width=Inches(4.6))
            except Exception:
                paragraph = self.document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(diagram_text)
                self._set_run_font(run, "Courier New", 11)

    def _set_run_font(self, run, font_name, size_pt):
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.font.bold = False
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            r_fonts.set(qn(attr), font_name)


def write_docx_file(docx_path, conversion_result):
    if not HAS_DOCX:
        raise RuntimeError("Exportacao DOCX indisponivel: instale 'python-docx'.")

    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(11)

    asset_lookup = {}
    for asset in conversion_result.diagram_assets:
        asset_lookup[asset.web_path.replace("\\", "/")] = asset.png_path
        asset_lookup[os.path.basename(asset.web_path)] = asset.png_path
        asset_lookup[os.path.basename(asset.png_path)] = asset.png_path

    with tempfile.TemporaryDirectory(prefix="pgn_docx_") as temp_dir:
        if conversion_result.summaries:
            title = document.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.add_run("Sumario")
            title_run.bold = True
            title_run.font.size = Pt(14)

            for summary in conversion_result.summaries:
                line = document.add_paragraph(style=None)
                parts = [f"{summary.index}. {summary.title}"]
                meta = " - ".join(
                    part for part in (summary.event, summary.site, summary.date, summary.result) if part
                )
                if meta:
                    parts.append(meta)
                line.add_run(" | ".join(parts))

            document.add_paragraph()

        for index, block in enumerate(conversion_result.blocks):
            parser = DocxBlockParser(document, asset_lookup, temp_dir, conversion_result.diagram_style)
            parser.feed(block)
            parser.close()

            if index != len(conversion_result.blocks) - 1:
                document.add_paragraph()

        footer = document.add_paragraph(
            f"Gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')}"
        )
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].italic = True

        document.save(docx_path)


# ====================== APLICATIVO TKINTER ======================

class LegacyTkApp:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("PGN → Livro de Xadrez (Vex. python-chess)")
        self.root.geometry("1450x920")
        self.root.configure(bg="#f4f4f4")
        self.project_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
        
        self.blocks = None
        self.conversion_result = None
        self.queue = queue.Queue()
        self.pgn_dir = "."
        self.preview_dir = None
        self.html_viewer_job = None
        self.html_highlight_job = None
        self.html_viewer_dirty = False
        self.viewer_last_mode = "completo"
        self.css_apply_job = None
        self.suspend_html_modified = False
        self.suspend_css_modified = False
        self.css_user_modified = False
        self.selected_game_indexes = None
        self.auto_viewer_var = tk.BooleanVar(value=True)
        self.html_renderer = self._create_html_renderer()
        self.diagram_style_options = dict(chess_diagrams.get_diagram_style_options())
        default_diagram_label = next(iter(self.diagram_style_options))
        self.diagram_style_label = tk.StringVar(value=default_diagram_label)
        self.css_preset_options = dict(get_css_preset_options())
        default_css_preset_label = next(iter(self.css_preset_options))
        self.css_preset_label = tk.StringVar(value=default_css_preset_label)
        self.exercise_mode_options = {
            "Livro completo": "book",
            "Somente exercícios": "exercises",
            "Livro + exercícios": "both",
        }
        self.exercise_mode_label = tk.StringVar(value="Livro completo")

        top = tk.Frame(self.root, bg="#2c3e50", pady=20)
        top.pack(fill="x")
        
        # Frame de butões
        btn_frame = tk.Frame(top, bg="#2c3e50")
        btn_frame.pack()

        btns = [
            ("Abrir PGN", self.abrir, "#3498db"),
            ("Validar PGN", self.validar_pgn, "#607d8b"),
            ("Selecionar Partidas", self.selecionar_partidas, "#546e7a"),
            ("Processar PGN", self.iniciar_processamento, "#27ae60"),
            ("Limpar Cache", self.limpar_cache_diagramas, "#795548"),
            ("Ver no Navegador", self.preview, "#e67e22"),
            ("SALVAR HTML", self.salvar_html, "#9b59b6"),
        ]
        if HAS_EPUB:
            btns.append(("Salvar EPUB", self.salvar_epub, "#1abc9c"))
        btns.append(("Salvar DOCX", self.salvar_docx, "#8e6e53"))
        btns.append(("Salvar PDF", self.salvar_pdf, "#455a64"))

        for texto, func, cor in btns:
            tk.Button(btn_frame, text=texto, command=func, bg=cor, fg="white",
                      font=("Arial", 11, "bold"), width=18, height=2).pack(side="left", padx=10)

        # Barra de Progresso
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(top, variable=self.progress_var, mode='indeterminate')
        self.progress_bar.pack(fill="x", padx=40, pady=(15, 0))
        
        self.status_lbl = tk.Label(top, text="Aguardando arquivo...", bg="#2c3e50", fg="#bdc3c7", font=("Arial", 10))
        self.status_lbl.pack(pady=5)

        options_frame = tk.Frame(top, bg="#2c3e50")
        options_frame.pack(pady=(4, 0))
        tk.Label(
            options_frame,
            text="Diagramas:",
            bg="#2c3e50",
            fg="#ecf0f1",
            font=("Arial", 10, "bold"),
        ).pack(side="left", padx=(0, 8))
        diagram_combo = ttk.Combobox(
            options_frame,
            textvariable=self.diagram_style_label,
            state="readonly",
            values=list(self.diagram_style_options.keys()),
            width=18,
        )
        diagram_combo.pack(side="left")
        diagram_combo.bind("<<ComboboxSelected>>", self.on_diagram_style_changed)
        tk.Label(
            options_frame,
            text="Tema:",
            bg="#2c3e50",
            fg="#ecf0f1",
            font=("Arial", 10, "bold"),
        ).pack(side="left", padx=(16, 8))
        css_preset_combo = ttk.Combobox(
            options_frame,
            textvariable=self.css_preset_label,
            state="readonly",
            values=list(self.css_preset_options.keys()),
            width=16,
        )
        css_preset_combo.pack(side="left")
        css_preset_combo.bind("<<ComboboxSelected>>", self.on_css_preset_changed)
        tk.Label(
            options_frame,
            text="Exercícios:",
            bg="#2c3e50",
            fg="#ecf0f1",
            font=("Arial", 10, "bold"),
        ).pack(side="left", padx=(16, 8))
        ttk.Combobox(
            options_frame,
            textvariable=self.exercise_mode_label,
            state="readonly",
            values=list(self.exercise_mode_options.keys()),
            width=16,
        ).pack(side="left")

        # Layout Paned
        pan = tk.PanedWindow(self.root, orient="horizontal", sashwidth=6)
        pan.pack(fill="both", expand=True, padx=15, pady=15)

        left = tk.Frame(pan, bg="white")
        self.left_tabs = ttk.Notebook(left)
        self.left_tabs.pack(fill="both", expand=True)

        pgn_tab = tk.Frame(self.left_tabs, bg="white")
        tk.Label(pgn_tab, text="Cole ou abra o PGN aqui", font=("Arial", 12, "bold"), bg="white").pack(anchor="w", pady=5)
        self.txt_pgn = scrolledtext.ScrolledText(pgn_tab, font=("Consolas", 11))
        self.txt_pgn.pack(fill="both", expand=True)
        self.left_tabs.add(pgn_tab, text="PGN")

        viewer_tab = tk.Frame(self.left_tabs, bg="white")
        viewer_top = tk.Frame(viewer_tab, bg="white")
        viewer_top.pack(fill="x", pady=5)
        tk.Label(viewer_top, text="HTML Viewer", font=("Arial", 12, "bold"), bg="white").pack(side="left")
        ttk.Checkbutton(
            viewer_top,
            text="Auto",
            variable=self.auto_viewer_var,
            command=self.on_auto_viewer_toggle,
        ).pack(side="right")
        tk.Button(
            viewer_top,
            text="Atualizar Viewer",
            command=self.refresh_html_viewer_now,
            bg="#4f81bd",
            fg="white",
            font=("Arial", 9, "bold"),
            width=15,
        ).pack(side="right", padx=(0, 8))
        self.viewer_status = tk.Label(
            viewer_tab,
            text="Preview aguardando HTML...",
            bg="white",
            fg="#666",
            font=("Arial", 10),
        )
        self.viewer_status.pack(anchor="w", pady=(0, 6))

        viewer_container = tk.Frame(viewer_tab, bg="#d9d9d9", highlightthickness=1, highlightbackground="#c8c8c8")
        viewer_container.pack(fill="both", expand=True)
        self.viewer_host = tk.Frame(viewer_container, bg="white")
        self.viewer_host.pack(fill="both", expand=True, padx=1, pady=1)
        self.viewer_host.bind("<Configure>", self.on_viewer_host_configure)
        self.left_tabs.add(viewer_tab, text="HTML Viewer")

        css_tab = tk.Frame(self.left_tabs, bg="white")
        css_top = tk.Frame(css_tab, bg="white")
        css_top.pack(fill="x", pady=(0, 6))
        tk.Label(css_top, text="CSS", font=("Arial", 12, "bold"), bg="white").pack(side="left")
        tk.Button(
            css_top,
            text="Carregar CSS",
            command=self.load_css_file,
            bg="#4f81bd",
            fg="white",
            font=("Arial", 9, "bold"),
            width=12,
        ).pack(side="right", padx=(6, 0))
        tk.Button(
            css_top,
            text="Salvar CSS",
            command=self.save_css_file,
            bg="#6d9e5b",
            fg="white",
            font=("Arial", 9, "bold"),
            width=12,
        ).pack(side="right")
        self.txt_css = scrolledtext.ScrolledText(css_tab, font=("Consolas", 10))
        self.txt_css.pack(fill="both", expand=True)
        self.txt_css.bind("<<Modified>>", self.on_css_modified)
        self.left_tabs.add(css_tab, text="CSS")
        self.left_tabs.bind("<<NotebookTabChanged>>", self.on_left_tab_changed)
        pan.add(left)

        right = tk.Frame(pan, bg="white")
        tk.Label(right, text="HTML gerado", font=("Arial", 12, "bold"), bg="white").pack(anchor="w", pady=5)
        self.txt_html = scrolledtext.ScrolledText(right, font=("Georgia", 12))
        self.txt_html.pack(fill="both", expand=True)
        self.txt_html.bind("<<Modified>>", self.on_html_modified)
        self._configure_html_editor_tags()
        pan.add(right)
        self.txt_html.edit_modified(False)
        self.set_css_text(self.build_selected_preset_css())
        if self.html_renderer is not None:
            try:
                self.html_renderer.attach_to_tk(
                    self.viewer_host,
                    line_click_callback=self.on_html_viewer_line_click,
                    load_finished_callback=self.on_html_viewer_load_finished,
                )
                self.start_qt_event_pump()
            except Exception as exc:
                self.html_renderer = None
                self.viewer_status.config(text=f"Viewer indisponível: {exc}")

    def get_selected_diagram_style(self):
        return self.diagram_style_options[self.diagram_style_label.get()]

    def get_selected_css_preset(self):
        return self.css_preset_options.get(self.css_preset_label.get(), "classic")

    def get_selected_exercise_mode(self):
        return self.exercise_mode_options.get(self.exercise_mode_label.get(), "book")

    def build_selected_preset_css(self):
        return load_css_preset(
            self.get_selected_css_preset(),
            diagram_style=self.get_selected_diagram_style(),
        )

    def get_effective_diagram_style(self):
        if self.conversion_result is not None:
            return self.conversion_result.diagram_style
        return self.get_selected_diagram_style()

    def _create_html_renderer(self):
        if not HAS_HTML_VIEWER:
            return None
        try:
            return HtmlPreviewRenderer(self.project_dir)
        except Exception:
            return None

    def start_qt_event_pump(self):
        if self.html_renderer is None:
            return

        def _pump():
            if self.html_renderer is None:
                return
            try:
                self.html_renderer.process_events()
            except Exception:
                return
            self.root.after(40, _pump)

        self.root.after(40, _pump)

    def on_viewer_host_configure(self, event=None):
        if self.html_renderer is None:
            return
        width = event.width if event is not None else self.viewer_host.winfo_width()
        height = event.height if event is not None else self.viewer_host.winfo_height()
        try:
            self.html_renderer.resize(width, height)
        except Exception:
            pass

    def on_html_viewer_line_click(self, line_number):
        self.root.after(0, lambda: self.highlight_html_line(line_number))

    def on_html_viewer_load_finished(self, ok):
        def _update():
            if not ok:
                self.viewer_status.config(text="Falha ao carregar o HTML no viewer.")
                return
            self.html_viewer_dirty = False
            mode_label = "parcial" if self.viewer_last_mode == "parcial" else "completo"
            self.viewer_status.config(
                text=f"Preview {mode_label} atualizado em {datetime.now().strftime('%H:%M:%S')}"
            )

        self.root.after(0, _update)

    def on_html_modified(self, event=None):
        if self.suspend_html_modified:
            self.txt_html.edit_modified(False)
            return
        if not self.txt_html.edit_modified():
            return
        self.txt_html.edit_modified(False)
        self.html_viewer_dirty = True
        self.schedule_html_highlight()
        if self.html_renderer is None:
            self.viewer_status.config(
                text=f"Viewer indisponível: {HTML_VIEWER_IMPORT_ERROR or 'erro de inicialização'}."
            )
            return
        self.update_viewer_pending_status()
        self.schedule_html_viewer_update()

    def on_css_modified(self, event=None):
        if self.suspend_css_modified:
            self.txt_css.edit_modified(False)
            return
        if not self.txt_css.edit_modified():
            return
        self.txt_css.edit_modified(False)
        self.css_user_modified = True
        self.schedule_css_apply()

    def on_left_tab_changed(self, event=None):
        if self.left_tabs.tab(self.left_tabs.select(), "text") == "HTML Viewer" and self.html_viewer_dirty:
            self.update_viewer_pending_status()
            self.schedule_html_viewer_update(delay_ms=250)

    def on_auto_viewer_toggle(self):
        self.update_viewer_pending_status()
        if self.html_viewer_job is not None:
            self.root.after_cancel(self.html_viewer_job)
            self.html_viewer_job = None
        if self.auto_viewer_var.get() and self.html_viewer_dirty and self.is_viewer_tab_active():
            self.schedule_html_viewer_update(delay_ms=250)

    def schedule_html_viewer_update(self, delay_ms=900):
        if self.html_renderer is None:
            return
        if not self.should_auto_refresh_viewer():
            return
        if self.html_viewer_job is not None:
            self.root.after_cancel(self.html_viewer_job)
        self.html_viewer_job = self.root.after(delay_ms, self.refresh_html_viewer)

    def refresh_html_viewer_now(self):
        if self.html_viewer_job is not None:
            self.root.after_cancel(self.html_viewer_job)
            self.html_viewer_job = None
        self.refresh_html_viewer()

    def schedule_html_highlight(self, delay_ms=120):
        if self.html_highlight_job is not None:
            self.root.after_cancel(self.html_highlight_job)
        self.html_highlight_job = self.root.after(delay_ms, self.apply_html_syntax_highlighting)

    def schedule_css_apply(self, delay_ms=300):
        if self.css_apply_job is not None:
            self.root.after_cancel(self.css_apply_job)
        self.css_apply_job = self.root.after(delay_ms, self.apply_css_changes)

    def _configure_html_editor_tags(self):
        self.html_syntax_tags = [
            "html_tag_bracket",
            "html_tag_name",
            "html_attr_name",
            "html_attr_value",
            "html_comment",
            "html_doctype",
            "html_css_selector",
            "html_css_property",
            "html_css_value",
        ]
        self.txt_html.tag_configure("viewer_line_highlight", background="#fff3a3")
        self.txt_html.tag_configure("html_tag_bracket", foreground="#9a3412")
        self.txt_html.tag_configure("html_tag_name", foreground="#1d4ed8")
        self.txt_html.tag_configure("html_attr_name", foreground="#b45309")
        self.txt_html.tag_configure("html_attr_value", foreground="#15803d")
        self.txt_html.tag_configure("html_comment", foreground="#6b7280")
        self.txt_html.tag_configure("html_doctype", foreground="#0f766e")
        self.txt_html.tag_configure("html_css_selector", foreground="#7c2d12")
        self.txt_html.tag_configure("html_css_property", foreground="#0369a1")
        self.txt_html.tag_configure("html_css_value", foreground="#047857")

    def _clear_html_syntax_tags(self):
        for tag_name in self.html_syntax_tags:
            self.txt_html.tag_remove(tag_name, "1.0", "end")

    def _add_html_tag_range(self, tag_name, start_offset, end_offset):
        if end_offset <= start_offset:
            return
        start_index = f"1.0 + {start_offset} chars"
        end_index = f"1.0 + {end_offset} chars"
        self.txt_html.tag_add(tag_name, start_index, end_index)

    def apply_html_syntax_highlighting(self):
        self.html_highlight_job = None
        text = self.txt_html.get("1.0", "end-1c")
        self._clear_html_syntax_tags()
        if not text:
            return

        for match in re.finditer(r"<!--.*?-->", text, re.DOTALL):
            self._add_html_tag_range("html_comment", match.start(), match.end())

        for match in re.finditer(r"<!DOCTYPE[^>]*>", text, re.IGNORECASE):
            self._add_html_tag_range("html_doctype", match.start(), match.end())

        tag_pattern = re.compile(r"<(/?)([A-Za-z][\w:-]*)([^<>]*?)(/?)>", re.DOTALL)
        attr_pattern = re.compile(
            r'([A-Za-z_:][\w:.-]*)(\s*=\s*)(".*?"|\'.*?\'|[^\s"\'>/=]+)',
            re.DOTALL,
        )

        for match in tag_pattern.finditer(text):
            full_start, full_end = match.span()
            slash, tag_name, attrs, self_close = match.groups()
            tag_name_start = full_start + 1 + len(slash)
            tag_name_end = tag_name_start + len(tag_name)
            attrs_start = tag_name_end

            self._add_html_tag_range("html_tag_bracket", full_start, full_start + 1 + len(slash))
            self._add_html_tag_range("html_tag_name", tag_name_start, tag_name_end)
            self._add_html_tag_range("html_tag_bracket", full_end - 1 - len(self_close), full_end)

            for attr_match in attr_pattern.finditer(attrs):
                name_start = attrs_start + attr_match.start(1)
                name_end = attrs_start + attr_match.end(1)
                value_start = attrs_start + attr_match.start(3)
                value_end = attrs_start + attr_match.end(3)
                self._add_html_tag_range("html_attr_name", name_start, name_end)
                self._add_html_tag_range("html_attr_value", value_start, value_end)

        style_block_pattern = re.compile(r"<style\b[^>]*>(.*?)</style>", re.IGNORECASE | re.DOTALL)
        css_rule_pattern = re.compile(r"([^{]+)\{([^}]*)\}", re.DOTALL)
        css_prop_pattern = re.compile(r"([A-Za-z-]+)\s*:\s*([^;}{]+)")

        for style_match in style_block_pattern.finditer(text):
            css_text = style_match.group(1)
            css_start = style_match.start(1)
            for rule_match in css_rule_pattern.finditer(css_text):
                selector_start = css_start + rule_match.start(1)
                selector_end = css_start + rule_match.end(1)
                self._add_html_tag_range("html_css_selector", selector_start, selector_end)

                body_text = rule_match.group(2)
                body_start = css_start + rule_match.start(2)
                for prop_match in css_prop_pattern.finditer(body_text):
                    prop_start = body_start + prop_match.start(1)
                    prop_end = body_start + prop_match.end(1)
                    value_start = body_start + prop_match.start(2)
                    value_end = body_start + prop_match.end(2)
                    self._add_html_tag_range("html_css_property", prop_start, prop_end)
                    self._add_html_tag_range("html_css_value", value_start, value_end)

    def get_current_css_text(self):
        diagram_style = self.get_effective_diagram_style()
        if not hasattr(self, "txt_css"):
            return load_css_preset(self.get_selected_css_preset(), diagram_style=diagram_style)
        css_text = self.txt_css.get("1.0", "end-1c")
        if not css_text.strip():
            return load_css_preset(self.get_selected_css_preset(), diagram_style=diagram_style)
        return ensure_diagram_font_css(css_text, diagram_style)

    def set_css_text(self, css_text):
        self.suspend_css_modified = True
        try:
            self.txt_css.delete("1.0", "end")
            self.txt_css.insert("1.0", css_text or "")
            self.txt_css.edit_modified(False)
        finally:
            self.suspend_css_modified = False

    def on_css_preset_changed(self, _event=None):
        if not hasattr(self, "txt_css"):
            return
        self.set_css_text(self.build_selected_preset_css())
        self.css_user_modified = self.get_selected_css_preset() != "classic"
        self.apply_css_changes()

    def on_diagram_style_changed(self, _event=None):
        if not hasattr(self, "txt_css") or self.css_user_modified:
            return
        self.set_css_text(self.build_selected_preset_css())

    def extract_css_from_html(self, html_text):
        match = re.search(r"<style\b[^>]*>(.*?)</style>", html_text, re.IGNORECASE | re.DOTALL)
        return match.group(1).strip() if match else ""

    def merge_css_into_html(self, html_text, css_text):
        css_text = css_text or ""
        style_block = f"<style>{css_text}</style>"

        if re.search(r"<style\b[^>]*>.*?</style>", html_text, re.IGNORECASE | re.DOTALL):
            return re.sub(
                r"<style\b[^>]*>.*?</style>",
                style_block,
                html_text,
                count=1,
                flags=re.IGNORECASE | re.DOTALL,
            )

        if re.search(r"</head>", html_text, re.IGNORECASE):
            return re.sub(r"</head>", style_block + "</head>", html_text, count=1, flags=re.IGNORECASE)

        return style_block + html_text

    def ensure_preview_dir(self):
        if self.preview_dir and os.path.isdir(self.preview_dir):
            return self.preview_dir
        self.preview_dir = tempfile.mkdtemp(prefix="pgn_preview_")
        return self.preview_dir

    def _externalize_html_css_for_viewer(self, html_text, css_href):
        link_tag = f'<link rel="stylesheet" href="{css_href}" data-preview-css="1">'
        html_without_inline_css = re.sub(
            r"<style\b[^>]*>.*?</style>",
            "",
            html_text,
            flags=re.IGNORECASE | re.DOTALL,
        )

        if re.search(r"<link\b[^>]*data-preview-css=['\"]1['\"][^>]*>", html_without_inline_css, re.IGNORECASE):
            return re.sub(
                r"<link\b[^>]*data-preview-css=['\"]1['\"][^>]*>",
                link_tag,
                html_without_inline_css,
                count=1,
                flags=re.IGNORECASE,
            )

        if re.search(r"<link\b[^>]*href=['\"][^'\"]*style\.css[^'\"]*['\"][^>]*>", html_without_inline_css, re.IGNORECASE):
            return re.sub(
                r"<link\b[^>]*href=['\"][^'\"]*style\.css[^'\"]*['\"][^>]*>",
                link_tag,
                html_without_inline_css,
                count=1,
                flags=re.IGNORECASE,
            )

        if re.search(r"</head>", html_without_inline_css, re.IGNORECASE):
            return re.sub(
                r"</head>",
                link_tag + "</head>",
                html_without_inline_css,
                count=1,
                flags=re.IGNORECASE,
            )

        return link_tag + html_without_inline_css

    def _prepare_preview_css(self, css_text):
        prepared_css = css_text
        for spec in chess_diagrams.get_diagram_font_specs():
            font_filename = spec["font_filename"]
            font_uri = QUrl.fromLocalFile(spec["font_path"]).toString()
            prepared_css = re.sub(
                rf'url\((["\']?)(?:[^)"\']*/)?{re.escape(font_filename)}\1\)',
                f'url("{font_uri}")',
                prepared_css,
                flags=re.IGNORECASE,
            )
        return prepared_css

    def prepare_viewer_html_bundle(self, html_text):
        preview_dir = self.ensure_preview_dir()
        css_text = self._prepare_preview_css(self.get_current_css_text())
        css_path = os.path.join(preview_dir, "style.css")
        html_path = os.path.join(preview_dir, "preview.html")

        with open(css_path, "w", encoding="utf-8") as file_obj:
            file_obj.write(css_text)

        self.copy_current_support_files(preview_dir)

        viewer_html = self._externalize_html_css_for_viewer(html_text, "style.css")
        return viewer_html, html_path

    def apply_css_changes(self):
        self.css_apply_job = None
        html_text = self.txt_html.get("1.0", "end-1c")
        if not html_text.strip():
            return

        merged_html = self.merge_css_into_html(html_text, self.get_current_css_text())
        if merged_html != html_text:
            yview = self.txt_html.yview()
            insert_index = self.txt_html.index("insert")
            self.suspend_html_modified = True
            try:
                self.txt_html.delete("1.0", "end")
                self.txt_html.insert("1.0", merged_html)
                self.txt_html.edit_modified(False)
            finally:
                self.suspend_html_modified = False
            try:
                self.txt_html.mark_set("insert", insert_index)
            except Exception:
                pass
            self.txt_html.yview_moveto(yview[0])
            self.schedule_html_highlight(delay_ms=10)

        self.html_viewer_dirty = True
        self.update_viewer_pending_status()
        self.schedule_html_viewer_update(delay_ms=400)

    def is_viewer_tab_active(self):
        return self.left_tabs.tab(self.left_tabs.select(), "text") == "HTML Viewer"

    def is_large_html_document(self):
        html_text = self.txt_html.get("1.0", "end-1c")
        total_lines = int(self.txt_html.index("end-1c").split(".")[0])
        return len(html_text) > 45000 or total_lines > 900

    def should_auto_refresh_viewer(self):
        return self.auto_viewer_var.get() and self.is_viewer_tab_active() and not self.is_large_html_document()

    def update_viewer_pending_status(self):
        if self.html_renderer is None:
            self.viewer_status.config(
                text=f"Viewer indisponível: {HTML_VIEWER_IMPORT_ERROR or 'erro de inicialização'}."
            )
        elif self.is_large_html_document():
            self.viewer_status.config(text="Preview pendente. HTML grande: 'Atualizar Viewer' carrega a secao atual.")
        elif not self.auto_viewer_var.get():
            self.viewer_status.config(text="Preview pendente. Auto desligado; use 'Atualizar Viewer'.")
        else:
            self.viewer_status.config(text="Preview pendente...")

    def get_html_insert_offset(self):
        try:
            return int(self.txt_html.count("1.0", "insert", "chars")[0])
        except Exception:
            return 0

    def build_viewer_html(self, html_text):
        if not self.is_large_html_document():
            return html_text, "completo"

        body_match = re.search(r"(<body\b[^>]*>)(.*?)(</body>)", html_text, re.IGNORECASE | re.DOTALL)
        if not body_match:
            return html_text, "completo"

        body_prefix = html_text[:body_match.start(2)]
        body_content = body_match.group(2)
        body_suffix = html_text[body_match.end(2):]
        hr_matches = list(re.finditer(r"<hr\b[^>]*>", body_content, re.IGNORECASE))
        if not hr_matches:
            return html_text, "completo"

        cursor_offset = max(0, self.get_html_insert_offset() - body_match.start(2))
        parts = []
        separators = []
        last_end = 0
        current_part_index = 0

        for idx, match in enumerate(hr_matches):
            parts.append(body_content[last_end:match.start()])
            separators.append(match.group(0))
            if match.start() <= cursor_offset:
                current_part_index = idx + 1
            last_end = match.end()
        parts.append(body_content[last_end:])

        current_part_index = max(0, min(current_part_index, len(parts) - 1))
        target_chars = 6000
        start_part = current_part_index
        end_part = current_part_index
        current_size = len(parts[current_part_index].strip())

        def left_addition_size(index):
            return len(separators[index]) + len(parts[index].strip())

        def right_addition_size(index):
            return len(separators[index - 1]) + len(parts[index].strip())

        while current_size < target_chars and (start_part > 0 or end_part < len(parts) - 1):
            left_candidate = None
            right_candidate = None

            if start_part > 0:
                left_size = left_addition_size(start_part - 1)
                left_candidate = (abs(target_chars - (current_size + left_size)), left_size, "left")

            if end_part < len(parts) - 1:
                right_size = right_addition_size(end_part + 1)
                right_candidate = (abs(target_chars - (current_size + right_size)), right_size, "right")

            candidates = [candidate for candidate in (left_candidate, right_candidate) if candidate is not None]
            if not candidates:
                break

            _, added_size, side = min(candidates, key=lambda item: (item[0], item[1]))
            current_size += added_size
            if side == "left":
                start_part -= 1
            else:
                end_part += 1

        selected_parts = []
        for idx in range(start_part, end_part + 1):
            if selected_parts and idx - 1 < len(separators):
                selected_parts.append(separators[idx - 1])
            selected_parts.append(parts[idx])

        partial_body = "".join(selected_parts).strip()
        if not partial_body:
            partial_body = parts[current_part_index]

        return body_prefix + partial_body + body_suffix, "parcial"

    def refresh_html_viewer(self):
        self.html_viewer_job = None
        if self.html_renderer is None:
            return

        html_text = self.txt_html.get("1.0", "end-1c").strip()
        if not html_text:
            self.html_renderer.clear()
            self.html_viewer_dirty = False
            self.viewer_status.config(text="Preview aguardando HTML...")
            return

        self.viewer_status.config(text="Atualizando preview...")

        try:
            viewer_html, viewer_mode = self.build_viewer_html(html_text)
            viewer_html, viewer_html_path = self.prepare_viewer_html_bundle(viewer_html)
            self.viewer_last_mode = viewer_mode
            self.html_renderer.load_html_document(
                viewer_html,
                viewer_html_path,
            )
        except Exception as exc:
            self.viewer_status.config(text=f"Falha no preview: {exc}")

    def set_html_text(self, html_text):
        css_text = self.extract_css_from_html(html_text)
        if css_text:
            normalized_css = ensure_diagram_font_css(
                css_text,
                self.get_effective_diagram_style(),
            )
            self.set_css_text(normalized_css)
            html_text = self.merge_css_into_html(html_text, normalized_css)
        self.suspend_html_modified = True
        try:
            self.txt_html.delete("1.0", "end")
            self.txt_html.insert("1.0", html_text)
            self.txt_html.edit_modified(False)
        finally:
            self.suspend_html_modified = False
        self.html_viewer_dirty = True
        self.schedule_html_highlight(delay_ms=10)
        if self.html_renderer is None:
            self.viewer_status.config(
                text=f"Viewer indisponível: {HTML_VIEWER_IMPORT_ERROR or 'erro de inicialização'}."
            )
        else:
            self.update_viewer_pending_status()
            self.schedule_html_viewer_update(delay_ms=350)

    def load_css_file(self):
        arq = filedialog.askopenfilename(filetypes=[("CSS", "*.css"), ("Todos", "*.*")])
        if not arq:
            return
        try:
            with open(arq, "r", encoding="utf-8", errors="ignore") as f:
                css_text = ensure_diagram_font_css(
                    f.read(),
                    self.get_effective_diagram_style(),
                )
            self.set_css_text(css_text)
            self.css_user_modified = True
            self.apply_css_changes()
            self.status_lbl.config(text=f"CSS carregado: {os.path.basename(arq)}")
        except Exception as exc:
            messagebox.showerror("Erro CSS", str(exc))

    def save_css_file(self):
        css_text = self.get_current_css_text()
        if not css_text.strip():
            messagebox.showwarning("Aviso", "Nao ha CSS para salvar.")
            return
        arq = filedialog.asksaveasfilename(defaultextension=".css", filetypes=[("CSS", "*.css")])
        if not arq:
            return
        try:
            with open(arq, "w", encoding="utf-8") as f:
                f.write(css_text)
            self.status_lbl.config(text=f"CSS salvo: {os.path.basename(arq)}")
        except Exception as exc:
            messagebox.showerror("Erro CSS", str(exc))

    def highlight_html_line(self, line_number):
        total_lines = int(self.txt_html.index("end-1c").split(".")[0])
        line_number = max(1, min(line_number, total_lines))

        self.txt_html.tag_remove("viewer_line_highlight", "1.0", "end")
        start = f"{line_number}.0"
        end = f"{line_number}.end"
        self.txt_html.tag_add("viewer_line_highlight", start, end)
        self.txt_html.mark_set("insert", start)
        self.txt_html.see(start)
        self.viewer_status.config(
            text=f"Preview atualizado em {datetime.now().strftime('%H:%M:%S')} | linha aproximada {line_number}"
        )

    def get_current_html_text(self):
        html_text = self.txt_html.get("1.0", "end-1c")
        return self.merge_css_into_html(html_text, self.get_current_css_text())

    def copy_current_support_files(self, output_dir):
        if self.conversion_result is None:
            return

        if chess_diagrams.uses_merida_style(self.conversion_result.diagram_style):
            chess_diagrams.copy_support_files(self.conversion_result.diagram_style, output_dir)

        if not self.conversion_result.diagram_assets:
            return

        diagram_dir = os.path.join(output_dir, "Diagrams")
        os.makedirs(diagram_dir, exist_ok=True)
        for asset in self.conversion_result.diagram_assets:
            for asset_path in (asset.svg_path, asset.png_path):
                if not asset_path or not os.path.isfile(asset_path):
                    continue
                target_path = os.path.join(diagram_dir, os.path.basename(asset_path))
                if os.path.abspath(asset_path) == os.path.abspath(target_path):
                    continue
                shutil.copyfile(asset_path, target_path)

    def write_current_html_bundle(self, html_path):
        html_text = self.get_current_html_text().strip()
        if not html_text:
            raise RuntimeError("Nao ha HTML para salvar.")

        output_dir = os.path.dirname(html_path)
        os.makedirs(output_dir, exist_ok=True)
        self.copy_current_support_files(output_dir)

        with open(html_path, "w", encoding="utf-8") as f:
            f.write(html_text)

        with open(os.path.join(output_dir, "style.css"), "w", encoding="utf-8") as f:
            f.write(self.get_current_css_text())

    def abrir(self):
        arq = filedialog.askopenfilename(filetypes=[("PGN", "*.pgn")])
        if arq:
            self.pgn_dir = os.path.dirname(arq)
            self.txt_pgn.delete("1.0", "end")
            try:
                with open(arq, "r", encoding="utf-8", errors="ignore") as f:
                    self.txt_pgn.insert("1.0", f.read())
                self.status_lbl.config(text=f"Arquivo carregado: {os.path.basename(arq)}")
            except Exception as e:
                messagebox.showerror("Erro", f"Não foi possível ler o arquivo: {e}")

    def validar_pgn(self):
        texto = self.txt_pgn.get("1.0", "end").strip()
        if not texto:
            messagebox.showwarning("Aviso", "Cole ou abra um arquivo PGN primeiro!")
            return

        issues = validate_pgn(texto)
        report = format_validation_report(issues)
        if issues:
            messagebox.showwarning("Validação PGN", report)
            error_count = sum(1 for issue in issues if issue.severity == "error")
            warning_count = len(issues) - error_count
            self.status_lbl.config(text=f"Validacao: {error_count} erro(s), {warning_count} aviso(s).")
        else:
            messagebox.showinfo("Validação PGN", report)
            self.status_lbl.config(text="Validacao: nenhum problema encontrado.")

    def selecionar_partidas(self):
        texto = self.txt_pgn.get("1.0", "end").strip()
        if not texto:
            messagebox.showwarning("Aviso", "Cole ou abra um arquivo PGN primeiro!")
            return

        summaries = scan_pgn_headers(texto)
        if not summaries:
            messagebox.showwarning("Selecionar Partidas", "Nenhuma partida PGN foi encontrada.")
            return

        window = tk.Toplevel(self.root)
        window.title("Selecionar Partidas")
        window.geometry("760x520")
        window.transient(self.root)
        window.grab_set()

        filters = tk.Frame(window)
        filters.pack(fill="x", padx=10, pady=8)
        player_var = tk.StringVar()
        eco_var = tk.StringVar()
        result_var = tk.StringVar()
        tk.Label(filters, text="Jogador").pack(side="left")
        tk.Entry(filters, textvariable=player_var, width=16).pack(side="left", padx=(4, 10))
        tk.Label(filters, text="ECO").pack(side="left")
        tk.Entry(filters, textvariable=eco_var, width=10).pack(side="left", padx=(4, 10))
        tk.Label(filters, text="Resultado").pack(side="left")
        tk.Entry(filters, textvariable=result_var, width=10).pack(side="left", padx=(4, 10))

        list_frame = tk.Frame(window)
        list_frame.pack(fill="both", expand=True, padx=10)
        canvas = tk.Canvas(list_frame)
        scrollbar = ttk.Scrollbar(list_frame, orient="vertical", command=canvas.yview)
        inner = tk.Frame(canvas)
        inner.bind("<Configure>", lambda _event: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=inner, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        selected = set(self.selected_game_indexes or [summary.index for summary in summaries])
        check_vars = {}

        def sync_visible_selection():
            for index, var in check_vars.items():
                if var.get():
                    selected.add(index)
                else:
                    selected.discard(index)

        def populate():
            sync_visible_selection()
            for child in inner.winfo_children():
                child.destroy()
            check_vars.clear()
            visible = filter_game_summaries(
                summaries,
                player=player_var.get(),
                eco=eco_var.get(),
                result=result_var.get(),
            )
            for summary in visible:
                var = tk.BooleanVar(value=summary.index in selected)
                check_vars[summary.index] = var
                label = (
                    f"{summary.index}. {summary.white} - {summary.black}"
                    f" [{summary.eco}]" if summary.eco else f"{summary.index}. {summary.white} - {summary.black}"
                )
                if summary.result:
                    label += f" | {summary.result}"
                if summary.event:
                    label += f" | {summary.event}"
                ttk.Checkbutton(inner, text=label, variable=var).pack(anchor="w", pady=2)

        for var in (player_var, eco_var, result_var):
            var.trace_add("write", lambda *_args: populate())

        actions = tk.Frame(window)
        actions.pack(fill="x", padx=10, pady=8)

        def accept():
            sync_visible_selection()
            self.selected_game_indexes = sorted(selected)
            total = len(summaries)
            chosen = len(self.selected_game_indexes)
            if chosen == total:
                self.selected_game_indexes = None
                self.status_lbl.config(text=f"Selecao: todas as {total} partida(s).")
            else:
                self.status_lbl.config(text=f"Selecao: {chosen}/{total} partida(s).")
            window.destroy()

        ttk.Button(actions, text="OK", command=accept).pack(side="right", padx=4)
        ttk.Button(actions, text="Cancelar", command=window.destroy).pack(side="right", padx=4)
        ttk.Button(actions, text="Marcar todas", command=lambda: (selected.update(s.index for s in summaries), populate())).pack(side="left", padx=4)
        ttk.Button(actions, text="Desmarcar todas", command=lambda: (selected.clear(), populate())).pack(side="left", padx=4)
        populate()
        self.root.wait_window(window)

    def limpar_cache_diagramas(self):
        diagram_dir = os.path.join(os.path.abspath(self.pgn_dir or "."), "Diagrams")
        removed = chess_diagrams.clear_diagram_cache(diagram_dir)
        messagebox.showinfo(
            "Cache de Diagramas",
            f"{removed} arquivo(s) de cache removido(s) em:\n{diagram_dir}",
        )
        self.status_lbl.config(text=f"Cache de diagramas limpo: {removed} arquivo(s).")

    def iniciar_processamento(self):
        texto = self.txt_pgn.get("1.0", "end").strip()
        if not texto:
            messagebox.showwarning("Aviso", "Cole ou abra um arquivo PGN primeiro!")
            return

        self.blocks = None
        self.conversion_result = None
        self.set_html_text("")
        self.progress_bar.start(10)
        self.status_lbl.config(text="Processando...")
        diagram_style = self.get_selected_diagram_style()
        
        # Thread para não travar a UI
        threading.Thread(
            target=processar_pgn_worker,
            args=(
                texto,
                self.queue,
                self.pgn_dir,
                diagram_style,
                self.selected_game_indexes,
                self.get_selected_exercise_mode(),
            ),
            daemon=True,
        ).start()
        self.root.after(100, self.check_queue)

    def check_queue(self):
        try:
            while True:
                task = self.queue.get_nowait()
                msg_type, data = task
                
                if msg_type == "status":
                    self.status_lbl.config(text=data)
                
                elif msg_type == "error":
                    self.progress_bar.stop()
                    self.status_lbl.config(text="Falha no processamento.")
                    messagebox.showerror("Erro no Processamento", data)
                    return
                
                elif msg_type == "done":
                    self.conversion_result = data
                    self.blocks = data.blocks
                    self.progress_bar.stop()
                    warning_suffix = f" ({len(data.warnings)} aviso(s))" if data.warnings else ""
                    self.status_lbl.config(
                        text=f"Pronto! {len(self.blocks)} partidas convertidas.{warning_suffix}"
                    )
                    
                    # Gerar HTML final
                    css_text = self.get_current_css_text() if self.css_user_modified else None
                    full_html = gerar_html_final(
                        self.blocks,
                        diagram_style=data.diagram_style,
                        css_text=css_text,
                        summaries=data.summaries,
                    )
                    self.set_html_text(full_html)
                    if data.warnings:
                        messagebox.showwarning(
                            "Processamento concluído com avisos",
                            "\n".join(data.warnings[:10]),
                        )
                    else:
                        messagebox.showinfo("Sucesso", "Processamento concluído com sucesso!")
                    return # Fim do polling
                    
        except queue.Empty:
            pass
        except Exception as exc:
            self.progress_bar.stop()
            self.status_lbl.config(text="Falha no processamento.")
            messagebox.showerror("Erro no Processamento", str(exc))
            return
        
        self.root.after(100, self.check_queue)

    def preview(self):
        if not self.txt_html.get("1.0", "end").strip():
            messagebox.showwarning("Aviso", "Processe o PGN primeiro.")
            return

        try:
            self.preview_dir = tempfile.mkdtemp(prefix="pgn_preview_")
            html_path = os.path.join(self.preview_dir, "preview.html")
            self.write_current_html_bundle(html_path)
            webbrowser.open("file://" + html_path)
        except Exception as exc:
            messagebox.showerror("Erro no Preview", str(exc))

    def salvar_html(self):
        if not self.txt_html.get("1.0", "end").strip():
            messagebox.showwarning("Aviso", "Processe o PGN primeiro.")
            return

        arq = filedialog.asksaveasfilename(defaultextension=".html", filetypes=[("HTML", "*.html")])
        if not arq:
            return

        try:
            self.write_current_html_bundle(arq)
            messagebox.showinfo("OK!", "HTML + CSS salvos com sucesso!")
        except Exception as exc:
            messagebox.showerror("Erro ao Salvar", str(exc))

    def salvar_epub(self):
        texto = self.txt_pgn.get("1.0", "end").strip()
        if not texto:
            messagebox.showwarning("Aviso", "Cole ou abra um arquivo PGN primeiro!")
            return

        arq = filedialog.asksaveasfilename(defaultextension=".epub")
        if not arq:
            return

        try:
            with tempfile.TemporaryDirectory(prefix="pgn_epub_") as temp_dir:
                result = convert_pgn(
                    texto,
                    output_dir=temp_dir,
                    diagram_style=self.get_selected_diagram_style(),
                    selected_game_indexes=self.selected_game_indexes,
                    exercise_mode=self.get_selected_exercise_mode(),
                )
                book = epub.EpubBook()
                book.set_title("Livro de Xadrez - PGN")
                book.set_language("pt")

                style = epub.EpubItem(
                    uid="style",
                    file_name="style/style.css",
                    media_type="text/css",
                    content=build_css(
                        diagram_style=result.diagram_style,
                        font_href=chess_diagrams.get_diagram_font_href(result.diagram_style),
                    ).encode(),
                )
                book.add_item(style)

                if chess_diagrams.uses_merida_style(result.diagram_style):
                    support_files = chess_diagrams.get_support_file_bytes(result.diagram_style)
                    for file_name, file_bytes in support_files.items():
                        book.add_item(
                            epub.EpubItem(
                                uid=f"style_{os.path.splitext(file_name)[0]}",
                                file_name=f"style/Fonts/{file_name}",
                                media_type=chess_diagrams.get_support_file_media_type(file_name),
                                content=file_bytes,
                            )
                        )

                caps = []
                for i, cont in enumerate(result.blocks, 1):
                    summary = result.summaries[i - 1] if i - 1 < len(result.summaries) else None
                    chapter_title = summary.title if summary else f"Partida {i}"
                    chapter = epub.EpubHtml(title=chapter_title, file_name=f"p{i}.xhtml")
                    chapter.content = (
                        '<html><head><link rel="stylesheet" href="style/style.css"/>'
                        f'</head><body>{cont}</body></html>'
                    ).encode()
                    book.add_item(chapter)
                    caps.append(chapter)

                added_assets = set()
                for asset in result.diagram_assets:
                    asset_files = [
                        (asset.svg_path, "image/svg+xml"),
                        (asset.png_path, "image/png"),
                    ]
                    for asset_path, media_type in asset_files:
                        file_name = os.path.join(
                            "Diagrams",
                            os.path.basename(asset_path),
                        ).replace("\\", "/")
                        if file_name in added_assets or not os.path.isfile(asset_path):
                            continue
                        added_assets.add(file_name)
                        with open(asset_path, "rb") as file_obj:
                            book.add_item(
                                epub.EpubItem(
                                    uid=file_name,
                                    file_name=file_name,
                                    media_type=media_type,
                                    content=file_obj.read(),
                                )
                            )

                book.toc = caps
                book.spine = ["nav"] + caps
                book.add_item(epub.EpubNcx())
                book.add_item(epub.EpubNav())

                epub.write_epub(arq, book)
                self.conversion_result = result
                self.blocks = result.blocks
            messagebox.showinfo("Sucesso", f"EPUB criado: {arq}")
            
        except Exception as e:
            messagebox.showerror("Erro EPUB", str(e))

    def salvar_docx(self):
        if not HAS_DOCX:
            python_exe = sys.executable or "python"
            messagebox.showerror(
                "DOCX indisponivel",
                "A exportacao DOCX exige o pacote 'python-docx' no mesmo Python que executa o programa.\n\n"
                f"Python em uso:\n{python_exe}\n\n"
                "Instale com:\n"
                f"\"{python_exe}\" -m pip install python-docx",
            )
            return

        texto = self.txt_pgn.get("1.0", "end").strip()
        if not texto:
            messagebox.showwarning("Aviso", "Cole ou abra um arquivo PGN primeiro!")
            return

        arq = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("DOCX", "*.docx")])
        if not arq:
            return

        try:
            pasta = os.path.dirname(arq)
            result = convert_pgn(
                texto,
                output_dir=pasta,
                diagram_style=self.get_selected_diagram_style(),
                selected_game_indexes=self.selected_game_indexes,
                exercise_mode=self.get_selected_exercise_mode(),
            )
            write_docx_file(arq, result)
            self.conversion_result = result
            self.blocks = result.blocks
            messagebox.showinfo("Sucesso", f"DOCX criado: {arq}")
        except Exception as exc:
            messagebox.showerror("Erro DOCX", str(exc))

    def salvar_pdf(self):
        texto = self.txt_pgn.get("1.0", "end").strip()
        if not texto:
            messagebox.showwarning("Aviso", "Cole ou abra um arquivo PGN primeiro!")
            return

        arq = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF", "*.pdf")])
        if not arq:
            return

        try:
            with tempfile.TemporaryDirectory(prefix="pgn_pdf_source_") as temp_dir:
                result = convert_pgn(
                    texto,
                    output_dir=temp_dir,
                    diagram_style=self.get_selected_diagram_style(),
                    selected_game_indexes=self.selected_game_indexes,
                    exercise_mode=self.get_selected_exercise_mode(),
                )
                css_text = self.get_current_css_text() if self.css_user_modified else None
                write_pdf_file(arq, result, css_text=css_text)
            messagebox.showinfo("Sucesso", f"PDF criado: {arq}")
        except Exception as exc:
            messagebox.showerror("Erro PDF", str(exc))

    def run(self):
        self.root.mainloop()


if __name__ == "__main__":
    from pgn_qt_window import App as QtApp

    sys.exit(QtApp(sys.modules[__name__]).run())
